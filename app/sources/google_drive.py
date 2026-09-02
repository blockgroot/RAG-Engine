"""Google Drive adapter for native Docs plus PDF and DOCX files."""

from __future__ import annotations

import logging
from datetime import datetime
from io import BytesIO

import httpx

from ..config.settings import GoogleSettings
from ..core.exceptions import ConfigurationError, SourceError
from .base import SourceAdapter, SourceDocument, SourceRef

logger = logging.getLogger(__name__)

_API_BASE = "https://www.googleapis.com/drive/v3"
_FOLDER_MIME = "application/vnd.google-apps.folder"
_DOC_MIME = "application/vnd.google-apps.document"
_PDF_MIME = "application/pdf"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_SUPPORTED_MIMES = {_DOC_MIME, _PDF_MIME, _DOCX_MIME}
# `lastModifyingUser` costs nothing here: it rides along in the files.list
# request we already make, which is why "top editors" needs no extra call.
_LIST_FIELDS = (
    "nextPageToken,files(id,name,mimeType,modifiedTime,trashed,parents,"
    "lastModifyingUser(displayName))"
)
_MAX_WALK_DEPTH = 20  # guard against pathological trees; Drive's own limit is ~100.


def _editor_name(file: dict) -> str | None:
    """Display name of whoever last modified the file, if Drive told us.

    Drive omits `lastModifyingUser` for some files (a service account edit, a
    deleted account, a shared drive with restricted metadata). None is the
    honest answer -- those rows simply do not appear in an editor chart, which
    is better than attributing them to someone.
    """
    user = file.get("lastModifyingUser") or {}
    name = (user.get("displayName") or "").strip()
    return name or None


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader  # lazy: keep this out of the module import cost

    reader = PdfReader(BytesIO(data))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    import docx  # lazy: keep this out of the module import cost

    document = docx.Document(BytesIO(data))
    return "\n\n".join(p.text for p in document.paragraphs)


def _file_uri(file_id: str, mime: str) -> str:
    if mime == _DOC_MIME:
        return _doc_uri(file_id)
    return f"https://drive.google.com/file/d/{file_id}/view"


def _parse_dt(value: str | None) -> datetime | None:
    """Parse Drive's RFC3339 ``modifiedTime`` value."""
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def _doc_uri(file_id: str) -> str:
    return f"https://docs.google.com/document/d/{file_id}/edit"


class GoogleDriveAdapter(SourceAdapter):
    """Fetches native Google Docs living under one Drive folder (recursively)."""

    def __init__(
        self,
        token: str,
        folder_id: str,
        *,
        timeout: float = 15.0,
        settings: GoogleSettings | None = None,
    ) -> None:
        """Build an adapter from a resolved access token and folder id."""
        if not token:
            raise ConfigurationError(
                "GoogleDriveAdapter requires a non-empty Google OAuth access token."
            )
        if not folder_id:
            raise ConfigurationError(
                "GoogleDriveAdapter requires a non-empty Drive folder id."
            )
        self._token = token
        self._folder_id = folder_id
        self._timeout = timeout
        limits = settings or GoogleSettings.from_env()
        self._max_walk_folders = max(1, limits.max_walk_folders)
        self._max_documents = max(1, limits.max_documents)

    # -- interface ---------------------------------------------------------

    def list_documents(self) -> list[SourceRef]:
        try:
            files = self._walk_folder(self._folder_id)
        except SourceError:
            raise
        except Exception as exc:
            raise SourceError(f"Google Drive files.list failed: {exc}", cause=exc) from exc

        return [
            SourceRef(
                external_id=file["id"],
                title=file.get("name", "Untitled"),
                last_modified=_parse_dt(file.get("modifiedTime")),
                source_uri=_file_uri(file["id"], file.get("mimeType", "")),
                last_editor=_editor_name(file),
            )
            for file in files
        ]

    def fetch_document(self, external_id: str) -> SourceDocument:
        try:
            meta = self._get_file_metadata(
                external_id,
                fields="name,mimeType,modifiedTime,lastModifyingUser(displayName)",
            )
            mime = meta.get("mimeType")
            if mime == _DOC_MIME:
                content = self._export_markdown(external_id)
            elif mime == _PDF_MIME:
                content = _extract_pdf_text(self._download_media(external_id))
            elif mime == _DOCX_MIME:
                content = _extract_docx_text(self._download_media(external_id))
            else:
                raise SourceError(f"Unsupported Google Drive file type: {mime!r}")
        except SourceError:
            raise
        except Exception as exc:
            raise SourceError(
                f"Google Drive fetch_document({external_id}) failed: {exc}", cause=exc
            ) from exc

        return SourceDocument(
            external_id=external_id,
            title=meta.get("name", "Untitled"),
            content=content,
            source_uri=_file_uri(external_id, mime or ""),
            last_modified=_parse_dt(meta.get("modifiedTime")),
            last_editor=_editor_name(meta),
        )

    def get_last_modified(self, external_id: str) -> datetime | None:
        try:
            meta = self._get_file_metadata(external_id, fields="modifiedTime")
        except SourceError:
            raise
        except Exception as exc:
            raise SourceError(
                f"Google Drive get_last_modified({external_id}) failed: {exc}", cause=exc
            ) from exc
        return _parse_dt(meta.get("modifiedTime"))

    # -- HTTP helpers --------------------------------------------------------

    def _headers(self) -> dict[str, str]:
        return {"Authorization": f"Bearer {self._token}"}

    def _list_children(self, folder_id: str, page_token: str | None = None) -> dict:
        params = {
            "q": f"'{folder_id}' in parents and trashed = false",
            "fields": _LIST_FIELDS,
            "supportsAllDrives": "true",
            "includeItemsFromAllDrives": "true",
            "pageSize": 100,
        }
        if page_token:
            params["pageToken"] = page_token
        try:
            response = httpx.get(
                f"{_API_BASE}/files",
                params=params,
                headers=self._headers(),
                timeout=self._timeout,
            )
            response.raise_for_status()
        except httpx.HTTPError as exc:
            raise SourceError(f"Google Drive files.list failed: {exc}", cause=exc) from exc
        return response.json()

    def _walk_folder(self, root_folder_id: str) -> list[dict]:
        """Breadth-first walk of one Drive folder tree, within configured bounds."""
        docs: list[dict] = []
        visited: set[str] = set()
        queue: list[tuple[str, int]] = [(root_folder_id, 0)]
        truncated: str | None = None

        while queue:
            if len(visited) >= self._max_walk_folders:
                truncated = (
                    f"folder limit reached ({self._max_walk_folders}); "
                    f"{len(queue)} subfolder(s) not visited"
                )
                break
            folder_id, depth = queue.pop(0)
            if folder_id in visited:
                continue
            visited.add(folder_id)
            if depth > _MAX_WALK_DEPTH:
                continue

            page_token: str | None = None
            while True:
                data = self._list_children(folder_id, page_token)
                for file in data.get("files", []):
                    mime = file.get("mimeType")
                    if mime == _FOLDER_MIME:
                        if file["id"] not in visited:
                            queue.append((file["id"], depth + 1))
                    elif mime in _SUPPORTED_MIMES:
                        docs.append(file)
                if len(docs) >= self._max_documents:
                    truncated = f"document limit reached ({self._max_documents})"
                    break
                page_token = data.get("nextPageToken")
                if not page_token:
                    break
            if truncated:
                break

        if truncated:
            logger.warning(
                "Google Drive walk of folder %s was truncated: %s. Raise "
                "GOOGLE_MAX_WALK_FOLDERS / GOOGLE_MAX_DOCUMENTS, or point the "
                "connection at a narrower folder.",
                root_folder_id,
                truncated,
            )

        return docs[: self._max_documents]

    def _export_markdown(self, file_id: str) -> str:
        try:
            response = httpx.get(
                f"{_API_BASE}/files/{file_id}/export",
                params={"mimeType": "text/markdown"},
                headers=self._headers(),
                timeout=self._timeout,
            )
            response.raise_for_status()
        except httpx.HTTPError as exc:
            raise SourceError(
                f"Google Drive files.export({file_id}) failed: {exc}", cause=exc
            ) from exc
        return response.text

    def _download_media(self, file_id: str) -> bytes:
        """Download raw bytes for a PDF or DOCX file."""
        try:
            response = httpx.get(
                f"{_API_BASE}/files/{file_id}",
                params={"alt": "media", "supportsAllDrives": "true"},
                headers=self._headers(),
                timeout=self._timeout,
            )
            response.raise_for_status()
        except httpx.HTTPError as exc:
            raise SourceError(
                f"Google Drive files.get(alt=media, {file_id}) failed: {exc}", cause=exc
            ) from exc
        return response.content

    def _get_file_metadata(self, file_id: str, *, fields: str) -> dict:
        try:
            response = httpx.get(
                f"{_API_BASE}/files/{file_id}",
                params={
                    "fields": fields,
                    "supportsAllDrives": "true",
                },
                headers=self._headers(),
                timeout=self._timeout,
            )
            response.raise_for_status()
        except httpx.HTTPError as exc:
            raise SourceError(
                f"Google Drive files.get({file_id}) failed: {exc}", cause=exc
            ) from exc
        return response.json()
